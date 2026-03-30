"""
Article Exporter - Export articles in multiple formats.

Supports:
- HTML (already generated in Stage 9)
- Markdown (converted from HTML)
- PDF (converted from HTML)
- JSON (structured data)
- CSV (flat table format)
- XLSX (Excel format with multiple sheets)
"""

import logging
import json
import csv
import os
import re
import time
import random
from typing import Dict, Any, Optional, List
from pathlib import Path
from datetime import datetime

logger = logging.getLogger(__name__)

# PDF service retry configuration
PDF_MAX_RETRIES = int(os.getenv("PDF_MAX_RETRIES", "3"))
PDF_BASE_DELAY = float(os.getenv("PDF_BASE_DELAY", "1.0"))
PDF_MAX_DELAY = float(os.getenv("PDF_MAX_DELAY", "10.0"))


class ArticleExporter:
    """Export articles in multiple formats."""

    @staticmethod
    def export_all(
        article: Dict[str, Any],
        html_content: str,
        output_dir: Path,
        formats: List[str] = ["html", "markdown", "pdf", "json", "csv", "xlsx"],
    ) -> Dict[str, str]:
        """
        Export article in all requested formats.

        Args:
            article: Article dictionary
            html_content: Rendered HTML content
            output_dir: Output directory
            formats: List of formats to export (html, markdown, pdf, json, csv, xlsx)

        Returns:
            Dictionary mapping format names to file paths
        """
        output_dir.mkdir(parents=True, exist_ok=True)
        exported_files = {}

        # Base filename from headline (strip HTML tags first)
        headline = article.get("Headline", "article")
        # Strip HTML tags and decode entities
        from html import unescape
        clean_headline = re.sub(r'<[^>]+>', '', headline)  # Remove HTML tags
        clean_headline = unescape(clean_headline)  # Decode HTML entities
        slug = ArticleExporter._generate_slug(clean_headline)
        base_name = slug or "article"

        # Export each format
        if "html" in formats:
            html_path = output_dir / f"{base_name}.html"
            html_path.write_text(html_content, encoding="utf-8")
            exported_files["html"] = str(html_path)
            logger.info(f"✅ Exported HTML: {html_path}")

        if "markdown" in formats:
            try:
                md_path = output_dir / f"{base_name}.md"
                md_content = ArticleExporter._html_to_markdown(html_content)
                md_path.write_text(md_content, encoding="utf-8")
                exported_files["markdown"] = str(md_path)
                logger.info(f"✅ Exported Markdown: {md_path}")
            except Exception as e:
                logger.warning(f"Markdown export failed: {e}")

        if "pdf" in formats:
            try:
                pdf_path = output_dir / f"{base_name}.pdf"
                ArticleExporter._html_to_pdf(html_content, pdf_path)
                exported_files["pdf"] = str(pdf_path)
                logger.info(f"✅ Exported PDF: {pdf_path}")
            except Exception as e:
                logger.warning(f"PDF export failed: {e}")

        if "json" in formats:
            json_path = output_dir / f"{base_name}.json"
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(article, f, indent=2, ensure_ascii=False)
            exported_files["json"] = str(json_path)
            logger.info(f"✅ Exported JSON: {json_path}")

        if "csv" in formats:
            try:
                csv_path = output_dir / f"{base_name}.csv"
                ArticleExporter._to_csv(article, csv_path)
                exported_files["csv"] = str(csv_path)
                logger.info(f"✅ Exported CSV: {csv_path}")
            except Exception as e:
                logger.warning(f"CSV export failed: {e}")

        if "xlsx" in formats:
            try:
                xlsx_path = output_dir / f"{base_name}.xlsx"
                ArticleExporter._to_xlsx(article, xlsx_path)
                exported_files["xlsx"] = str(xlsx_path)
                logger.info(f"✅ Exported XLSX: {xlsx_path}")
            except Exception as e:
                logger.warning(f"XLSX export failed: {e}")

        if "docx" in formats:
            try:
                docx_path = output_dir / f"{base_name}.docx"
                ArticleExporter._export_docx(article, docx_path)
                exported_files["docx"] = str(docx_path)
                logger.info(f"✅ Exported DOCX: {docx_path}")
            except Exception as e:
                logger.warning(f"DOCX export failed: {e}")

        return exported_files

    @staticmethod
    def _html_to_markdown(html_content: str) -> str:
        """
        Convert HTML to Markdown using markdownify library.

        Args:
            html_content: HTML content

        Returns:
            Markdown content
        """
        from markdownify import markdownify as md

        # Extract body content if full HTML document
        body_match = re.search(r'<body[^>]*>(.*?)</body>', html_content, re.DOTALL | re.IGNORECASE)
        if body_match:
            html_content = body_match.group(1)

        # Remove script and style tags completely before markdown conversion
        # (markdownify's strip parameter doesn't reliably remove them)
        html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)

        # Convert to markdown
        result = md(html_content, heading_style="ATX")

        # Clean up extra whitespace
        result = re.sub(r'\n{3,}', '\n\n', result)

        return result.strip()

    @staticmethod
    def _html_to_pdf(html_content: str, output_path: Path) -> None:
        """
        Convert HTML to PDF with embedded images.

        Uses external PDF service with retry logic for transient failures.

        Args:
            html_content: HTML content
            output_path: Output PDF path
        """
        # Try to use external PDF service if available
        pdf_service_url = os.getenv(
            "PDF_SERVICE_URL",
            "https://clients--pdf-generation-fastapi-app.modal.run"
        )

        import requests
        import base64

        # Convert local image references to base64 data URLs for PDF generation
        html_with_images = ArticleExporter._embed_images_for_pdf(html_content, output_path.parent)

        # Add CSS for PDF margins
        html_with_images = ArticleExporter._add_pdf_margins(html_with_images)

        payload = {
            "html": html_with_images,
            "format": "A4",
            "landscape": False,
            "print_background": True,
            "prefer_css_page_size": True,
            "viewport_width": 1200,
            "device_scale_factor": 2,
            "color_scheme": "light",
            "margin": {
                "top": "25mm",
                "right": "20mm",
                "bottom": "25mm",
                "left": "20mm"
            }
        }

        last_error = None
        for attempt in range(PDF_MAX_RETRIES + 1):
            try:
                with requests.post(
                    f"{pdf_service_url}/convert",
                    headers={"Content-Type": "application/json"},
                    json=payload,
                    timeout=120,
                ) as response:
                    response.raise_for_status()
                    result = response.json()
                    pdf_data = base64.b64decode(result["pdf_base64"])
                    output_path.write_bytes(pdf_data)
                    logger.info(f"PDF generated via service ({len(pdf_data)} bytes)")
                    return

            except requests.exceptions.Timeout as e:
                last_error = e
                logger.warning(f"PDF service timeout (attempt {attempt + 1}/{PDF_MAX_RETRIES + 1})")
            except requests.exceptions.HTTPError as e:
                last_error = e
                # Check if retryable (5xx errors)
                if hasattr(e, 'response') and e.response is not None:
                    status_code = e.response.status_code
                    # Log response body for debugging (truncate for safety)
                    try:
                        response_body = e.response.text[:500] if e.response.text else "(empty)"
                    except Exception:
                        response_body = "(could not read response)"
                    if status_code < 500:
                        # Client error, don't retry - log full details
                        logger.error(f"PDF service client error ({status_code}): {e}")
                        logger.error(f"PDF service response body: {response_body}")
                        raise
                    logger.warning(f"PDF service server error ({status_code}): {response_body[:200]}")
                logger.warning(f"PDF service error (attempt {attempt + 1}/{PDF_MAX_RETRIES + 1}): {e}")
            except requests.exceptions.ConnectionError as e:
                last_error = e
                logger.warning(f"PDF service connection error (attempt {attempt + 1}/{PDF_MAX_RETRIES + 1}): {e}")
            except Exception as e:
                # Non-retryable error
                logger.error(f"PDF service failed: {e}")
                raise

            # Exponential backoff with jitter
            if attempt < PDF_MAX_RETRIES:
                delay = min(PDF_BASE_DELAY * (2 ** attempt), PDF_MAX_DELAY)
                jitter = random.uniform(0, delay * 0.1)
                time.sleep(delay + jitter)

        # All retries exhausted
        logger.warning(f"PDF service unavailable after {PDF_MAX_RETRIES + 1} attempts")
        raise last_error or Exception("PDF service failed")

    @staticmethod
    def _embed_images_for_pdf(html_content: str, base_path: Path) -> str:
        """
        Convert local image references to base64 data URLs for PDF generation.
        
        Args:
            html_content: HTML content with local image paths
            base_path: Base directory to resolve relative paths
            
        Returns:
            HTML with base64 data URLs
        """
        import re
        
        def replace_image_src(match):
            img_tag = match.group(0)
            src = match.group(1)
            
            # Skip if already a data URL or external URL
            if src.startswith(('data:', 'http://', 'https://')):
                return img_tag
            
            # Resolve relative path
            img_path = base_path / src
            if not img_path.exists():
                # Try resolving from project root
                project_root = Path(__file__).parent.parent.parent
                img_path = project_root / src
                if not img_path.exists():
                    logger.debug(f"Image not found for PDF: {src}")
                    return img_tag
            
            try:
                # Read and encode image
                img_data = img_path.read_bytes()

                # Determine MIME type
                ext = img_path.suffix.lower()
                mime_type = {
                    '.png': 'image/png',
                    '.jpg': 'image/jpeg',
                    '.jpeg': 'image/jpeg',
                    '.webp': 'image/webp',
                    '.gif': 'image/gif'
                }.get(ext, 'image/png')

                # Create base64 data URL
                b64_data = base64.b64encode(img_data).decode('utf-8')
                data_url = f"data:{mime_type};base64,{b64_data}"

                # Replace src in the img tag
                new_img_tag = img_tag.replace(f'src="{src}"', f'src="{data_url}"')

                logger.debug(f"Embedded image for PDF: {src} ({len(img_data)} bytes)")
                return new_img_tag

            except (OSError, IOError) as e:
                # File system errors (permission denied, file truncated, disk read error)
                logger.warning(f"Failed to read image file {src}: {type(e).__name__}: {e}")
                return img_tag
            except Exception as e:
                # Unexpected errors (encoding issues, etc.)
                logger.warning(f"Failed to embed image {src} for PDF: {type(e).__name__}: {e}")
                return img_tag
        
        # Find and replace all image src attributes
        pattern = r'<img[^>]+src="([^"]+)"[^>]*>'
        modified_html = re.sub(pattern, replace_image_src, html_content)
        
        return modified_html

    @staticmethod
    def _add_pdf_margins(html_content: str) -> str:
        """
        Add CSS margins for PDF export.
        
        Args:
            html_content: HTML content
            
        Returns:
            HTML with PDF margin CSS added
        """
        import re
        
        # CSS for PDF margins
        pdf_margin_css = """
        @page {
            margin: 25mm 20mm;
        }
        body {
            padding: 20px;
        }
        """
        
        # Insert CSS before closing </head> tag
        if '</head>' in html_content:
            html_content = html_content.replace('</head>', f'<style>{pdf_margin_css}</style></head>')
        elif '<style>' in html_content:
            # Insert before first </style> tag
            html_content = re.sub(r'(</style>)', f'{pdf_margin_css}\\1', html_content, count=1)
        else:
            # Insert before </head> or at start of <body>
            if '</head>' in html_content:
                html_content = html_content.replace('</head>', f'<style>{pdf_margin_css}</style></head>')
            elif '<body>' in html_content:
                html_content = html_content.replace('<body>', f'<body><style>{pdf_margin_css}</style>')
        
        return html_content

    @staticmethod
    def _to_csv(article: Dict[str, Any], output_path: Path) -> None:
        """
        Export article to CSV format (flat table).

        Args:
            article: Article dictionary
            output_path: Output CSV path
        """
        rows = []

        # Main article fields
        rows.append(["Field", "Value"])
        rows.append(["Headline", article.get("Headline", "")])
        rows.append(["Subtitle", article.get("Subtitle", "")])
        rows.append(["Meta Title", article.get("Meta_Title", "")])
        rows.append(["Meta Description", article.get("Meta_Description", "")])
        rows.append(["Teaser", article.get("Teaser", "")])
        rows.append(["Direct Answer", article.get("Direct_Answer", "")])
        # Convert HTML to single-line for easier Excel viewing
        intro = article.get("Intro", "")
        rows.append(["Intro", ArticleExporter._html_to_single_line(intro) if intro else ""])
        rows.append(["Word Count", article.get("word_count", 0)])
        rows.append(["Read Time", article.get("read_time", 0)])

        # Sections
        rows.append([])
        rows.append(["Section", "Title", "Content"])
        for i in range(1, 10):
            title = article.get(f"section_{i:02d}_title", "")
            content = article.get(f"section_{i:02d}_content", "")
            if title or content:
                # Convert HTML to single-line for easier Excel viewing
                content_single_line = ArticleExporter._html_to_single_line(content)
                rows.append([f"Section {i}", title, content_single_line])  # Full content, no truncation

        # FAQ
        rows.append([])
        rows.append(["FAQ", "Question", "Answer"])
        for i in range(1, 7):
            question = article.get(f"faq_{i:02d}_question", "")
            answer = article.get(f"faq_{i:02d}_answer", "")
            if question:
                rows.append([f"FAQ {i}", question, answer])

        # PAA
        rows.append([])
        rows.append(["PAA", "Question", "Answer"])
        for i in range(1, 5):
            question = article.get(f"paa_{i:02d}_question", "")
            answer = article.get(f"paa_{i:02d}_answer", "")
            if question:
                rows.append([f"PAA {i}", question, answer])

        # Write CSV
        with open(output_path, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerows(rows)

    @staticmethod
    def _to_xlsx(article: Dict[str, Any], output_path: Path) -> None:
        """
        Export article to XLSX format (Excel with multiple sheets).

        Args:
            article: Article dictionary
            output_path: Output XLSX path
        """
        try:
            import openpyxl
            from openpyxl.styles import Font, Alignment

            wb = openpyxl.Workbook()

            # Sheet 1: Overview
            ws1 = wb.active
            ws1.title = "Overview"
            ws1.append(["Field", "Value"])
            ws1.append(["Headline", article.get("Headline", "")])
            ws1.append(["Subtitle", article.get("Subtitle", "")])
            ws1.append(["Meta Title", article.get("Meta_Title", "")])
            ws1.append(["Meta Description", article.get("Meta_Description", "")])
            # Convert HTML fields to single-line
            intro = article.get("Intro", "")
            direct_answer = article.get("Direct_Answer", "")
            ws1.append(["Intro", ArticleExporter._html_to_single_line(intro) if intro else ""])
            ws1.append(["Direct Answer", ArticleExporter._html_to_single_line(direct_answer) if direct_answer else ""])
            ws1.append(["Word Count", article.get("word_count", 0)])
            ws1.append(["Read Time", article.get("read_time", 0)])
            ws1.append(["Publication Date", article.get("publication_date", "")])

            # Sheet 2: Sections
            ws2 = wb.create_sheet("Sections")
            ws2.append(["Section", "Title", "Content"])
            for i in range(1, 10):
                title = article.get(f"section_{i:02d}_title", "")
                content = article.get(f"section_{i:02d}_content", "")
                if title or content:
                    # Convert HTML to single-line for easier Excel viewing
                    content_single_line = ArticleExporter._html_to_single_line(content)
                    ws2.append([f"Section {i}", title, content_single_line])

            # Sheet 3: FAQ
            ws3 = wb.create_sheet("FAQ")
            ws3.append(["#", "Question", "Answer"])
            for i in range(1, 7):
                question = article.get(f"faq_{i:02d}_question", "")
                answer = article.get(f"faq_{i:02d}_answer", "")
                if question:
                    ws3.append([i, question, answer])

            # Sheet 4: PAA
            ws4 = wb.create_sheet("PAA")
            ws4.append(["#", "Question", "Answer"])
            for i in range(1, 5):
                question = article.get(f"paa_{i:02d}_question", "")
                answer = article.get(f"paa_{i:02d}_answer", "")
                if question:
                    ws4.append([i, question, answer])

            # Format headers
            for ws in [ws1, ws2, ws3, ws4]:
                header_font = Font(bold=True)
                for cell in ws[1]:
                    cell.font = header_font
                    cell.alignment = Alignment(horizontal="left", vertical="top")

            wb.save(output_path)

        except ImportError:
            logger.warning("openpyxl not installed, skipping XLSX export. Install with: pip install openpyxl")
            raise

    @staticmethod
    def _export_docx(article: Dict[str, Any], output_path: Path, language: str = "de") -> None:
        """
        Export article to a professionally formatted DOCX for lawyer review.

        Produces a Word document with styled headings, colored highlight boxes,
        proper paragraph spacing, and bold/italic preserved from HTML content.
        """
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml
        from html import unescape

        def strip_html(text: str) -> str:
            """Strip HTML tags and decode entities."""
            if not text:
                return ""
            text = re.sub(r'<br\s*/?>', '\n', text)
            text = re.sub(r'<li[^>]*>', '\n• ', text)
            text = re.sub(r'</?(ul|ol|div|section|aside|article|nav|header|footer)[^>]*>', '', text)
            text = re.sub(r'<[^>]+>', '', text)
            text = unescape(text)
            text = re.sub(r'\n{3,}', '\n\n', text)
            return text.strip()

        def add_rich_paragraphs(doc, html_content: str):
            """Parse HTML content and add paragraphs with bold/italic formatting."""
            if not html_content:
                return
            # Split on block-level tags
            blocks = re.split(r'<(?:p|div|li)[^>]*>|</(?:p|div|li)>', html_content)
            for block in blocks:
                block = block.strip()
                if not block or block in ('', '<ul>', '</ul>', '<ol>', '</ol>'):
                    continue
                # Check if it was a list item
                is_bullet = False
                if block.startswith('• ') or '<li' in block:
                    is_bullet = True
                    block = re.sub(r'^• ', '', block)

                # Extract bold/italic runs
                p = doc.add_paragraph(style='List Bullet' if is_bullet else 'Normal')
                p.paragraph_format.space_after = Pt(6)
                # Parse inline formatting
                parts = re.split(r'(<strong>|</strong>|<b>|</b>|<em>|</em>|<i>|</i>)', block)
                bold = False
                italic = False
                for part in parts:
                    if part in ('<strong>', '<b>'):
                        bold = True
                        continue
                    elif part in ('</strong>', '</b>'):
                        bold = False
                        continue
                    elif part in ('<em>', '<i>'):
                        italic = True
                        continue
                    elif part in ('</em>', '</i>'):
                        italic = False
                        continue
                    # Strip remaining tags
                    clean = re.sub(r'<[^>]+>', '', part)
                    clean = unescape(clean).strip()
                    if clean:
                        run = p.add_run(clean + ' ')
                        run.font.size = Pt(11)
                        run.bold = bold
                        run.italic = italic

        def add_shaded_box(doc, label: str, content: str, shade_color: str = "E8F4E8"):
            """Add a shaded box with a label and content (like the HTML callout boxes)."""
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            # Add shading to paragraph
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shade_color}" w:val="clear"/>')
            p.paragraph_format.element.get_or_add_pPr().append(shading)
            # Left border
            pPr = p.paragraph_format.element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'  <w:left w:val="single" w:sz="24" w:space="8" w:color="1A5276"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            # Indent
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(label.upper())
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(26, 82, 118)

            p2 = doc.add_paragraph()
            shading2 = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shade_color}" w:val="clear"/>')
            p2.paragraph_format.element.get_or_add_pPr().append(shading2)
            p2.paragraph_format.left_indent = Cm(0.5)
            p2.paragraph_format.space_after = Pt(12)
            run2 = p2.add_run(strip_html(content))
            run2.font.size = Pt(11)

        def add_takeaway_box(doc, label: str, items: list, shade_color: str = "FFF8E1"):
            """Add a styled takeaway box with bullet points."""
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shade_color}" w:val="clear"/>')
            p.paragraph_format.element.get_or_add_pPr().append(shading)
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(f"⚡ {label}")
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(180, 95, 6)

            for item in items:
                bp = doc.add_paragraph()
                shading_b = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shade_color}" w:val="clear"/>')
                bp.paragraph_format.element.get_or_add_pPr().append(shading_b)
                bp.paragraph_format.left_indent = Cm(1.0)
                bp.paragraph_format.space_after = Pt(4)
                run = bp.add_run(f"→  {item}")
                run.font.size = Pt(11)

            # Final spacer
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)

        doc = Document()

        # ---- Style configuration ----
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.15

        # Heading styles
        for level in range(1, 4):
            h_style = doc.styles[f'Heading {level}']
            h_style.font.name = 'Calibri'
            h_style.font.color.rgb = RGBColor(26, 82, 118)
            if level == 1:
                h_style.font.size = Pt(22)
                h_style.paragraph_format.space_before = Pt(0)
                h_style.paragraph_format.space_after = Pt(6)
            elif level == 2:
                h_style.font.size = Pt(16)
                h_style.paragraph_format.space_before = Pt(18)
                h_style.paragraph_format.space_after = Pt(8)
            else:
                h_style.font.size = Pt(13)

        # Page margins
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # ---- Document content ----

        # 1. Title
        headline = strip_html(article.get("Headline", ""))
        if headline:
            doc.add_heading(headline, level=1)

        # 2. Subtitle
        subtitle = strip_html(article.get("Subtitle", ""))
        if subtitle:
            p = doc.add_paragraph()
            run = p.add_run(subtitle)
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(100, 100, 100)
            p.paragraph_format.space_after = Pt(16)

        # 3. Direct Answer (shaded box)
        direct_answer = article.get("Direct_Answer", "")
        if direct_answer:
            label = "Kurze Antwort" if language == "de" else "Quick Answer"
            add_shaded_box(doc, label, direct_answer)

        # 4. Key Takeaways (yellow box)
        takeaways = []
        for i in range(1, 4):
            t = article.get(f"key_takeaway_{i:02d}", "")
            if t:
                takeaways.append(strip_html(t))
        if takeaways:
            label = "Das Thema kurz und kompakt" if language == "de" else "Key Points at a Glance"
            add_takeaway_box(doc, label, takeaways)

        # 5. Intro
        intro = article.get("Intro", "")
        if intro:
            add_rich_paragraphs(doc, intro)

        # 6. Sections
        for i in range(1, 10):
            title = strip_html(article.get(f"section_{i:02d}_title", ""))
            content = article.get(f"section_{i:02d}_content", "")
            if not title and not content:
                continue
            if title:
                doc.add_heading(title, level=2)
            if content:
                add_rich_paragraphs(doc, content)

        # 7. FAQs
        faq_label = "Häufig gestellte Fragen" if language == "de" else "FAQ"
        has_faqs = False
        for i in range(1, 7):
            q = strip_html(article.get(f"faq_{i:02d}_question", ""))
            a = article.get(f"faq_{i:02d}_answer", "")
            if q and a:
                if not has_faqs:
                    doc.add_heading(faq_label, level=2)
                    has_faqs = True
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                run = p.add_run(q)
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(26, 82, 118)
                add_rich_paragraphs(doc, a)

        # 8. Legal disclaimer (gray box)
        disclaimer = strip_html(article.get("legal_disclaimer", ""))
        if disclaimer:
            label = "Rechtlicher Hinweis" if language == "de" else "Legal Notice"
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(24)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
            p.paragraph_format.element.get_or_add_pPr().append(shading)
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(f"⚖ {label}: ")
            run.bold = True
            run.font.size = Pt(10)
            run = p.add_run(disclaimer)
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)

        # 9. Sources
        sources = article.get("Sources", "")
        if sources:
            label = "Quellen" if language == "de" else "Sources"
            doc.add_heading(label, level=2)
            source_list = sources if isinstance(sources, list) else [
                l.strip() for l in sources.strip().split('\n') if l.strip()
            ]
            for src in source_list:
                p = doc.add_paragraph(str(src), style='List Number')
                p.runs[0].font.size = Pt(9) if p.runs else None
                p.runs[0].font.color.rgb = RGBColor(80, 80, 80) if p.runs else None

        doc.save(str(output_path))

    @staticmethod
    def _html_to_single_line(html_content: str) -> str:
        """
        Convert multi-line HTML to single-line for easier Excel viewing.
        
        Args:
            html_content: Multi-line HTML content
            
        Returns:
            Single-line HTML content
        """
        if not html_content:
            return ""
        
        # Remove newlines and extra whitespace between tags
        # Keep single space between tags and text
        single_line = re.sub(r'\s+', ' ', html_content)  # Replace all whitespace with single space
        single_line = re.sub(r'>\s+<', '><', single_line)  # Remove spaces between tags
        single_line = single_line.strip()
        
        return single_line
    
    @staticmethod
    def _generate_slug(text: str) -> str:
        """Generate URL slug from text."""
        if not text:
            return ""
        slug = text.lower()
        slug = re.sub(r'[^\w\s-]', '', slug)
        slug = re.sub(r'[-\s]+', '-', slug)
        return slug.strip('-')

